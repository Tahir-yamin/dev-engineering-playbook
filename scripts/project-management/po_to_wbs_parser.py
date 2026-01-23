"""
Purchase Order / Scope of Work to WBS Parser
=============================================
Converts PO documents and scope descriptions into WBS structure
for automatic schedule creation in MS Project and Primavera P6.

Requirements:
    pip install python-docx PyPDF2 openpyxl

Usage:
    from po_to_wbs_parser import POParser
    parser = POParser()
    wbs = parser.parse_text(scope_text)
"""

import re
import json
import csv
from typing import List, Dict, Optional
from datetime import datetime
import os

# Optional imports for document parsing
try:
    from docx import Document
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False

try:
    import PyPDF2
    HAS_PDF = True
except ImportError:
    HAS_PDF = False

try:
    import openpyxl
    HAS_EXCEL = True
except ImportError:
    HAS_EXCEL = False


class POParser:
    """
    Parse Purchase Orders and Scope of Work documents to generate WBS.
    """
    
    def __init__(self):
        """Initialize parser with standard patterns."""
        self.wbs_counter = {}  # Track WBS numbering
        
        # Common scope sections to look for
        self.section_keywords = [
            "scope of work", "scope of supply", "deliverables",
            "work packages", "tasks", "activities", "milestones",
            "phases", "requirements", "specifications",
            "engineering", "procurement", "construction",
            "installation", "commissioning", "testing",
            "design", "fabrication", "delivery"
        ]
        
        # Duration keywords for estimation
        self.duration_keywords = {
            "immediately": 1,
            "urgent": 3,
            "asap": 5,
            "week": 5,
            "weeks": 5,
            "month": 20,
            "months": 20,
            "days": 1,
            "within": 5
        }
        
    # ==========================================
    # DOCUMENT PARSING
    # ==========================================
    
    def parse_file(self, filepath: str) -> Dict:
        """
        Parse document file (Word, PDF, Excel, Text).
        
        Args:
            filepath: Path to document
            
        Returns:
            Dict with parsed content and generated WBS
        """
        ext = os.path.splitext(filepath)[1].lower()
        
        if ext in ['.docx', '.doc']:
            text = self._read_docx(filepath)
        elif ext == '.pdf':
            text = self._read_pdf(filepath)
        elif ext in ['.xlsx', '.xls']:
            return self._parse_excel(filepath)
        elif ext == '.txt':
            with open(filepath, 'r', encoding='utf-8') as f:
                text = f.read()
        else:
            raise ValueError(f"Unsupported file type: {ext}")
            
        return self.parse_text(text)
        
    def _read_docx(self, filepath: str) -> str:
        """Read Word document."""
        if not HAS_DOCX:
            raise ImportError("python-docx required: pip install python-docx")
            
        doc = Document(filepath)
        text_parts = []
        
        for para in doc.paragraphs:
            text_parts.append(para.text)
            
        # Also get tables
        for table in doc.tables:
            for row in table.rows:
                row_text = [cell.text for cell in row.cells]
                text_parts.append(" | ".join(row_text))
                
        return "\n".join(text_parts)
        
    def _read_pdf(self, filepath: str) -> str:
        """Read PDF document."""
        if not HAS_PDF:
            raise ImportError("PyPDF2 required: pip install PyPDF2")
            
        text_parts = []
        
        with open(filepath, 'rb') as f:
            reader = PyPDF2.PdfReader(f)
            for page in reader.pages:
                text_parts.append(page.extract_text())
                
        return "\n".join(text_parts)
        
    def _parse_excel(self, filepath: str) -> Dict:
        """Parse Excel file with structured WBS data."""
        if not HAS_EXCEL:
            raise ImportError("openpyxl required: pip install openpyxl")
            
        wb = openpyxl.load_workbook(filepath)
        ws = wb.active
        
        wbs_items = []
        headers = [cell.value for cell in ws[1]]
        
        # Find relevant columns
        wbs_col = self._find_column(headers, ['wbs', 'code', 'id'])
        name_col = self._find_column(headers, ['name', 'description', 'task', 'activity'])
        dur_col = self._find_column(headers, ['duration', 'days', 'weeks'])
        level_col = self._find_column(headers, ['level', 'indent'])
        
        for row in ws.iter_rows(min_row=2, values_only=True):
            if not any(row):
                continue
                
            item = {
                "wbs": row[wbs_col] if wbs_col is not None else self._auto_wbs(len(wbs_items)),
                "name": row[name_col] if name_col is not None else str(row[0]),
                "duration": self._parse_duration(row[dur_col]) if dur_col is not None else "5d",
                "level": int(row[level_col]) if level_col is not None else 1
            }
            
            if item["name"]:
                wbs_items.append(item)
                
        return {
            "source": filepath,
            "wbs": wbs_items,
            "metadata": {"parsed_from": "excel", "row_count": len(wbs_items)}
        }
        
    def _find_column(self, headers: List, keywords: List[str]) -> Optional[int]:
        """Find column index by keyword match."""
        for i, header in enumerate(headers):
            if header and any(kw in str(header).lower() for kw in keywords):
                return i
        return None
        
    # ==========================================
    # TEXT PARSING
    # ==========================================
    
    def parse_text(self, text: str) -> Dict:
        """
        Parse scope of work text and generate WBS.
        
        Args:
            text: Scope of work text content
            
        Returns:
            Dict with WBS structure and metadata
        """
        # Clean text
        text = self._clean_text(text)
        
        # Extract sections
        sections = self._extract_sections(text)
        
        # Generate WBS from sections
        wbs = self._generate_wbs(sections)
        
        # Extract project metadata
        metadata = self._extract_metadata(text)
        
        return {
            "source": "text_input",
            "wbs": wbs,
            "sections": sections,
            "metadata": metadata
        }
        
    def _clean_text(self, text: str) -> str:
        """Clean and normalize text."""
        # Remove excessive whitespace
        text = re.sub(r'\s+', ' ', text)
        # Normalize bullet points
        text = re.sub(r'[•●○►▪▸]', '-', text)
        return text.strip()
        
    def _extract_sections(self, text: str) -> List[Dict]:
        """Extract major sections from text."""
        sections = []
        
        # Pattern for numbered sections (1.0, 2.0, etc.)
        numbered_pattern = r'(\d+(?:\.\d+)*)\s*[.:\-]?\s*([A-Za-z][^0-9\n]+)'
        
        # Pattern for lettered sections (A., B., etc.)
        lettered_pattern = r'([A-Z])[.:\)]\s*([A-Za-z][^0-9\n]+)'
        
        # Pattern for bullet points
        bullet_pattern = r'[-*]\s+([A-Za-z][^\n]+)'
        
        # Find numbered sections first
        for match in re.finditer(numbered_pattern, text):
            sections.append({
                "code": match.group(1),
                "name": match.group(2).strip(),
                "type": "numbered"
            })
            
        # If no numbered sections, try lettered
        if not sections:
            for match in re.finditer(lettered_pattern, text):
                sections.append({
                    "code": match.group(1),
                    "name": match.group(2).strip(),
                    "type": "lettered"
                })
                
        # If still nothing, try bullets
        if not sections:
            for i, match in enumerate(re.finditer(bullet_pattern, text)):
                sections.append({
                    "code": str(i + 1),
                    "name": match.group(1).strip(),
                    "type": "bullet"
                })
                
        # If still nothing, use keyword-based extraction
        if not sections:
            sections = self._extract_by_keywords(text)
            
        return sections
        
    def _extract_by_keywords(self, text: str) -> List[Dict]:
        """Extract sections based on common scope keywords."""
        sections = []
        text_lower = text.lower()
        
        for keyword in self.section_keywords:
            if keyword in text_lower:
                # Find the sentence containing this keyword
                pattern = rf'[^.]*{keyword}[^.]*\.'
                matches = re.findall(pattern, text_lower)
                for match in matches:
                    sections.append({
                        "code": str(len(sections) + 1),
                        "name": keyword.title(),
                        "type": "keyword",
                        "context": match
                    })
                    
        return sections
        
    def _generate_wbs(self, sections: List[Dict]) -> List[Dict]:
        """Generate WBS structure from sections."""
        wbs = []
        level_1_counter = 0
        
        for section in sections:
            # Determine level from code
            if '.' in str(section.get('code', '')):
                parts = str(section['code']).split('.')
                level = len(parts)
            else:
                level = 1
                level_1_counter += 1
                
            # Create WBS code
            if level == 1:
                wbs_code = f"{level_1_counter}.0"
            else:
                wbs_code = section['code']
                
            # Estimate duration
            duration = self._estimate_duration(section['name'])
            
            wbs.append({
                "wbs": wbs_code,
                "name": section['name'][:100],  # Limit name length
                "level": level,
                "duration": f"{duration}d",
                "original_code": section.get('code')
            })
            
        return wbs
        
    def _estimate_duration(self, name: str) -> int:
        """Estimate duration based on activity name."""
        name_lower = name.lower()
        
        # Check for duration keywords
        for keyword, days in self.duration_keywords.items():
            if keyword in name_lower:
                return days
                
        # Default estimations based on activity type
        if any(kw in name_lower for kw in ['design', 'engineering', 'plan']):
            return 10
        elif any(kw in name_lower for kw in ['review', 'approve', 'sign']):
            return 3
        elif any(kw in name_lower for kw in ['procure', 'order', 'purchase']):
            return 15
        elif any(kw in name_lower for kw in ['fabricat', 'manufactur', 'construct']):
            return 20
        elif any(kw in name_lower for kw in ['install', 'erect', 'assemble']):
            return 15
        elif any(kw in name_lower for kw in ['test', 'commission', 'startup']):
            return 10
        elif any(kw in name_lower for kw in ['deliver', 'ship', 'transport']):
            return 5
        elif any(kw in name_lower for kw in ['mobiliz', 'demobiliz']):
            return 3
        else:
            return 5  # Default
            
    def _extract_metadata(self, text: str) -> Dict:
        """Extract project metadata from text."""
        metadata = {}
        
        # PO Number
        po_patterns = [
            r'P[O/.]?\s*[#:]?\s*(\d+[-/\w]*)',
            r'Purchase Order[:\s]+(\d+[-/\w]*)',
            r'Order\s+(?:No\.?|Number)[:\s]+(\d+[-/\w]*)'
        ]
        for pattern in po_patterns:
            match = re.search(pattern, text, re.IGNORECASE)
            if match:
                metadata['po_number'] = match.group(1)
                break
                
        # Contract Number
        contract_patterns = [
            r'Contract\s+(?:No\.?|Number)[:\s]+(\w+[-/\w]*)',
            r'Contract[:\s]+(\w+[-/\w]*)'
        ]
        for pattern in contract_patterns:
            match = re.search(pattern, text, re.IGNORECASE)
            if match:
                metadata['contract_number'] = match.group(1)
                break
                
        # Project Name
        project_patterns = [
            r'Project\s+(?:Name|Title)[:\s]+([^\n]+)',
            r'Subject[:\s]+([^\n]+)'
        ]
        for pattern in project_patterns:
            match = re.search(pattern, text, re.IGNORECASE)
            if match:
                metadata['project_name'] = match.group(1).strip()
                break
                
        # Dates
        date_patterns = [
            r'Start\s+Date[:\s]+(\d{1,2}[-/]\d{1,2}[-/]\d{2,4})',
            r'Completion\s+Date[:\s]+(\d{1,2}[-/]\d{1,2}[-/]\d{2,4})',
            r'Due\s+Date[:\s]+(\d{1,2}[-/]\d{1,2}[-/]\d{2,4})'
        ]
        for pattern in date_patterns:
            match = re.search(pattern, text, re.IGNORECASE)
            if match:
                key = 'start_date' if 'start' in pattern.lower() else 'end_date'
                metadata[key] = match.group(1)
                
        # Value
        value_patterns = [
            r'(?:Total|Contract|PO)\s+(?:Value|Amount|Price)[:\s]+\$?([\d,]+(?:\.\d{2})?)',
            r'\$([\d,]+(?:\.\d{2})?)'
        ]
        for pattern in value_patterns:
            match = re.search(pattern, text, re.IGNORECASE)
            if match:
                metadata['value'] = match.group(1)
                break
                
        return metadata
        
    def _auto_wbs(self, index: int) -> str:
        """Generate automatic WBS code."""
        return f"{(index // 10) + 1}.{(index % 10) + 1}"
        
    def _parse_duration(self, value) -> str:
        """Parse duration from various formats."""
        if not value:
            return "5d"
            
        value_str = str(value).lower()
        
        # Extract numbers
        numbers = re.findall(r'\d+', value_str)
        if not numbers:
            return "5d"
            
        num = int(numbers[0])
        
        # Determine unit
        if 'week' in value_str:
            return f"{num * 5}d"
        elif 'month' in value_str:
            return f"{num * 20}d"
        else:
            return f"{num}d"
            
    # ==========================================
    # OUTPUT GENERATION
    # ==========================================
    
    def to_ms_project_format(self, wbs_data: List[Dict]) -> List[Dict]:
        """Convert to MS Project automation format."""
        return wbs_data  # Already in correct format
        
    def to_p6_format(self, wbs_data: List[Dict]) -> List[Dict]:
        """Convert to P6 import format."""
        p6_format = []
        
        for item in wbs_data:
            # Generate activity ID (P6 style)
            activity_id = f"A{item['wbs'].replace('.', '')}"
            
            p6_format.append({
                "activity_id": activity_id,
                "name": item['name'],
                "duration": int(item['duration'].replace('d', '')) if 'd' in item['duration'] else 5,
                "wbs": item['wbs']
            })
            
        return p6_format
        
    def export_to_csv(self, wbs_data: List[Dict], output_path: str):
        """Export WBS to CSV for manual import."""
        if not wbs_data:
            print("No WBS data to export")
            return
            
        with open(output_path, 'w', newline='', encoding='utf-8') as f:
            writer = csv.DictWriter(f, fieldnames=wbs_data[0].keys())
            writer.writeheader()
            writer.writerows(wbs_data)
            
        print(f"✓ Exported WBS to {output_path}")
        
    def export_to_json(self, parsed_data: Dict, output_path: str):
        """Export full parsed data to JSON."""
        with open(output_path, 'w', encoding='utf-8') as f:
            json.dump(parsed_data, f, indent=2, default=str)
            
        print(f"✓ Exported to {output_path}")


class StandardWBSTemplates:
    """
    Standard WBS templates for common project types.
    """
    
    @staticmethod
    def epc_project() -> List[Dict]:
        """EPC (Engineering, Procurement, Construction) template."""
        return [
            {"wbs": "1.0", "name": "Project Management", "level": 1, "duration": "0d"},
            {"wbs": "1.1", "name": "Project Initiation", "level": 2, "duration": "5d"},
            {"wbs": "1.2", "name": "Project Planning", "level": 2, "duration": "10d"},
            {"wbs": "1.3", "name": "Project Monitoring & Control", "level": 2, "duration": "0d"},
            {"wbs": "1.4", "name": "Project Closeout", "level": 2, "duration": "5d"},
            
            {"wbs": "2.0", "name": "Engineering", "level": 1, "duration": "0d"},
            {"wbs": "2.1", "name": "Basic Engineering", "level": 2, "duration": "20d"},
            {"wbs": "2.2", "name": "Detailed Engineering", "level": 2, "duration": "30d"},
            {"wbs": "2.3", "name": "Engineering Review & Approval", "level": 2, "duration": "10d"},
            
            {"wbs": "3.0", "name": "Procurement", "level": 1, "duration": "0d"},
            {"wbs": "3.1", "name": "Vendor Qualification", "level": 2, "duration": "10d"},
            {"wbs": "3.2", "name": "Bid & Award", "level": 2, "duration": "20d"},
            {"wbs": "3.3", "name": "Material Delivery", "level": 2, "duration": "30d"},
            
            {"wbs": "4.0", "name": "Construction", "level": 1, "duration": "0d"},
            {"wbs": "4.1", "name": "Site Mobilization", "level": 2, "duration": "5d"},
            {"wbs": "4.2", "name": "Civil Works", "level": 2, "duration": "30d"},
            {"wbs": "4.3", "name": "Structural Works", "level": 2, "duration": "45d"},
            {"wbs": "4.4", "name": "Mechanical Installation", "level": 2, "duration": "30d"},
            {"wbs": "4.5", "name": "Electrical Installation", "level": 2, "duration": "20d"},
            {"wbs": "4.6", "name": "Instrumentation", "level": 2, "duration": "15d"},
            
            {"wbs": "5.0", "name": "Commissioning & Startup", "level": 1, "duration": "0d"},
            {"wbs": "5.1", "name": "Pre-Commissioning", "level": 2, "duration": "10d"},
            {"wbs": "5.2", "name": "Commissioning", "level": 2, "duration": "15d"},
            {"wbs": "5.3", "name": "Performance Testing", "level": 2, "duration": "10d"},
            {"wbs": "5.4", "name": "Handover", "level": 2, "duration": "5d"},
        ]
        
    @staticmethod
    def software_project() -> List[Dict]:
        """Software development project template."""
        return [
            {"wbs": "1.0", "name": "Initiation", "level": 1, "duration": "0d"},
            {"wbs": "1.1", "name": "Requirements Gathering", "level": 2, "duration": "10d"},
            {"wbs": "1.2", "name": "Feasibility Study", "level": 2, "duration": "5d"},
            {"wbs": "1.3", "name": "Project Charter", "level": 2, "duration": "2d"},
            
            {"wbs": "2.0", "name": "Design", "level": 1, "duration": "0d"},
            {"wbs": "2.1", "name": "System Architecture", "level": 2, "duration": "10d"},
            {"wbs": "2.2", "name": "UI/UX Design", "level": 2, "duration": "15d"},
            {"wbs": "2.3", "name": "Database Design", "level": 2, "duration": "5d"},
            {"wbs": "2.4", "name": "Design Review", "level": 2, "duration": "3d"},
            
            {"wbs": "3.0", "name": "Development", "level": 1, "duration": "0d"},
            {"wbs": "3.1", "name": "Backend Development", "level": 2, "duration": "30d"},
            {"wbs": "3.2", "name": "Frontend Development", "level": 2, "duration": "25d"},
            {"wbs": "3.3", "name": "Integration", "level": 2, "duration": "10d"},
            
            {"wbs": "4.0", "name": "Testing", "level": 1, "duration": "0d"},
            {"wbs": "4.1", "name": "Unit Testing", "level": 2, "duration": "10d"},
            {"wbs": "4.2", "name": "Integration Testing", "level": 2, "duration": "10d"},
            {"wbs": "4.3", "name": "UAT", "level": 2, "duration": "10d"},
            {"wbs": "4.4", "name": "Bug Fixes", "level": 2, "duration": "10d"},
            
            {"wbs": "5.0", "name": "Deployment", "level": 1, "duration": "0d"},
            {"wbs": "5.1", "name": "Staging Deployment", "level": 2, "duration": "3d"},
            {"wbs": "5.2", "name": "Production Deployment", "level": 2, "duration": "2d"},
            {"wbs": "5.3", "name": "Documentation", "level": 2, "duration": "5d"},
            {"wbs": "5.4", "name": "Training", "level": 2, "duration": "5d"},
        ]
        
    @staticmethod
    def procurement_project() -> List[Dict]:
        """Equipment procurement project template."""
        return [
            {"wbs": "1.0", "name": "Pre-Award", "level": 1, "duration": "0d"},
            {"wbs": "1.1", "name": "Scope Definition", "level": 2, "duration": "5d"},
            {"wbs": "1.2", "name": "Specifications", "level": 2, "duration": "10d"},
            {"wbs": "1.3", "name": "RFQ Preparation", "level": 2, "duration": "5d"},
            {"wbs": "1.4", "name": "Vendor Selection", "level": 2, "duration": "10d"},
            
            {"wbs": "2.0", "name": "Award", "level": 1, "duration": "0d"},
            {"wbs": "2.1", "name": "Contract Negotiation", "level": 2, "duration": "10d"},
            {"wbs": "2.2", "name": "PO Issuance", "level": 2, "duration": "2d"},
            {"wbs": "2.3", "name": "Kickoff Meeting", "level": 2, "duration": "1d"},
            
            {"wbs": "3.0", "name": "Manufacturing", "level": 1, "duration": "0d"},
            {"wbs": "3.1", "name": "Design & Engineering", "level": 2, "duration": "20d"},
            {"wbs": "3.2", "name": "Material Procurement", "level": 2, "duration": "30d"},
            {"wbs": "3.3", "name": "Fabrication", "level": 2, "duration": "45d"},
            {"wbs": "3.4", "name": "Factory Testing", "level": 2, "duration": "10d"},
            
            {"wbs": "4.0", "name": "Delivery", "level": 1, "duration": "0d"},
            {"wbs": "4.1", "name": "Packing & Shipping", "level": 2, "duration": "5d"},
            {"wbs": "4.2", "name": "Transportation", "level": 2, "duration": "15d"},
            {"wbs": "4.3", "name": "Site Receipt", "level": 2, "duration": "2d"},
            
            {"wbs": "5.0", "name": "Installation & Commissioning", "level": 1, "duration": "0d"},
            {"wbs": "5.1", "name": "Installation", "level": 2, "duration": "15d"},
            {"wbs": "5.2", "name": "Commissioning", "level": 2, "duration": "10d"},
            {"wbs": "5.3", "name": "Training", "level": 2, "duration": "5d"},
            {"wbs": "5.4", "name": "Handover", "level": 2, "duration": "2d"},
        ]


# ==========================================
# EXAMPLE USAGE
# ==========================================

if __name__ == "__main__":
    print("="*60)
    print("PO / Scope of Work to WBS Parser")
    print("="*60)
    
    # Example: Parse scope text
    sample_scope = """
    Purchase Order: PO-2026-001
    Project: Pump Station Upgrade
    Contract Value: $500,000
    
    SCOPE OF WORK:
    
    1.0 Engineering
        1.1 Process Design
        1.2 Mechanical Design
        1.3 Electrical Design
        
    2.0 Procurement
        2.1 Pumps and Motors
        2.2 Piping Materials
        2.3 Electrical Equipment
        
    3.0 Construction
        3.1 Civil Works
        3.2 Mechanical Installation
        3.3 Electrical Installation
        
    4.0 Commissioning
        4.1 Pre-Commissioning
        4.2 Performance Testing
        4.3 Handover
    """
    
    parser = POParser()
    result = parser.parse_text(sample_scope)
    
    print("\n📋 Extracted Metadata:")
    for key, value in result['metadata'].items():
        print(f"  {key}: {value}")
        
    print("\n📊 Generated WBS:")
    for item in result['wbs']:
        indent = "  " * (item['level'] - 1)
        print(f"  {indent}{item['wbs']} {item['name']} ({item['duration']})")
        
    # Export
    parser.export_to_csv(result['wbs'], "D:/wbs_output.csv")
    
    print("\n" + "="*60)
    print("Standard WBS Templates Available:")
    print("  - StandardWBSTemplates.epc_project()")
    print("  - StandardWBSTemplates.software_project()")
    print("  - StandardWBSTemplates.procurement_project()")
    print("="*60)
