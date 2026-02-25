import docx
import sys
import os

def extract_text(docx_path):
    if not os.path.exists(docx_path):
        print(f"Error: File not found at {docx_path}")
        return

    try:
        doc = docx.Document(docx_path)
        full_text = []
        for para in doc.paragraphs:
            if para.text.strip():
                full_text.append(para.text)
        
        # Also extract text from tables if any
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                     for para in cell.paragraphs:
                        if para.text.strip():
                            full_text.append(para.text)

        print("\n".join(full_text))
    except Exception as e:
        print(f"Error reading docx: {e}")

if __name__ == "__main__":
    if len(sys.argv) < 2:
        print("Usage: python extract_docx.py <path_to_docx>")
    else:
        extract_text(sys.argv[1])
