import json
import os
from datetime import datetime
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

BASE_DIR = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
PROFILE_PATH = os.path.join(BASE_DIR, "data", "master_profile.json")
OUTPUT_DIR = os.path.join(BASE_DIR, "generated")

def load_profile():
    with open(PROFILE_PATH, 'r', encoding='utf-8') as f:
        return json.load(f)

def generate_cover_letter(company_name, job_title, job_ref=None):
    profile = load_profile()
    p = profile["personal"]
    
    doc = Document()
    
    # Styles
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)

    # 1. Header (Contact Info)
    header = doc.add_paragraph()
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    name_run = header.add_run(p["name"].upper() + "\n")
    name_run.bold = True
    name_run.font.size = Pt(14)
    
    contact_info = f"{p['location']} | {p['phone']} | {p['email_primary']}\n{p['linkedin']}"
    header.add_run(contact_info).font.size = Pt(10)
    
    doc.add_paragraph() # Spacer

    # 2. Date and Recipient
    doc.add_paragraph(datetime.now().strftime("%B %d, %Y"))
    
    doc.add_paragraph(f"Hiring Manager\n{company_name}\nDoha, Qatar") # Fixed location based on JD context

    # 3. Salutation
    doc.add_paragraph(f"Re: Application for {job_title} Position" + (f" (Ref: {job_ref})" if job_ref else ""))
    
    doc.add_paragraph("Dear Hiring Manager,")

    # 4. Opening Hook (The "Why Me" immediately)
    opening = (
        f"With over 15 years of experience in Project Controls and Planning for major EPC and Oil & Gas "
        f"projects (up to USD 750M+), I am writing to express my strong interest in the {job_title} "
        f"role at {company_name}. My background in managing complex brownfield and greenfield portfolios, "
        f"combined with my expertise in Primavera P6, Schedule Risk Analysis, and FIDIC-based claims management, "
        f"aligns directly with the requirements of your projects team in Qatar."
    )
    doc.add_paragraph(opening)

    # 5. The "Value Add" (Bullet points matching JD)
    doc.add_paragraph("Throughout my career, I have consistently delivered results by optimizing schedules, "
                      "controlling critical paths, and leading planning teams. Key highlights relevant to this role include:")
    
    bullets = [
        "**Strategic Planning Leadership:** Led project controls for a USD 750M+ EPC portfolio, establishing Level-3/4 baselines and governing WBS/CBS structures across engineering, procurement, and construction phases.",
        "**Delay Analysis & Recovery:** Successfully substantiated EOT claims and variation orders using rigorous delay analysis (TIA/IAB) techniques under FIDIC contracts, protecting commercial interests.",
        "**Mentorship & Continuous Improvement:** Managed and mentored teams of up to 15 planning engineers, driving standardisation using Primavera P6 and implementing Power BI dashboards for executive reporting.",
        "**Resource & Cost Integration:** Expert in resource loading, levelling, and integrating cost with schedule (EVM) to provide accurate SPI/CPI forecasting and trend analysis."
    ]
    
    for b in bullets:
        para = doc.add_paragraph(style='List Bullet')
        parts = b.split("**")
        if len(parts) > 2:
            run = para.add_run(parts[1])
            run.bold = True
            para.add_run(parts[2])
        else:
            para.add_run(b)

    # 6. Conclusion & Call to Action
    closing = (
        f"I am eager to bring my proactive approach to {company_name}, ensuring that your FEED and execution "
        f"projects meet their strict man-hour, cost, and schedule targets. I am available for an immediate "
        f"interview and am fully open to relocation to Qatar."
    )
    doc.add_paragraph(closing)
    
    doc.add_paragraph("Sincerely,")
    doc.add_paragraph(p["name"])
    
    # Save
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    filename = f"Cover_Letter_TahirYamin_{company_name.replace(' ', '_')}_{timestamp}.docx"
    output_path = os.path.join(OUTPUT_DIR, filename)
    doc.save(output_path)
    
    print(f"Generated Cover Letter: {output_path}")

if __name__ == "__main__":
    generate_cover_letter("Wood", "Principal Planner", "26819")
