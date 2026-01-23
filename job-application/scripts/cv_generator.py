"""
ATS-Friendly CV Generator
=========================
Generates tailored CVs for specific job descriptions.
Optimizes keywords for ATS matching.

Usage:
    python cv_generator.py --job "job_description.txt"
    python cv_generator.py --keywords "Primavera P6,EPC,Claims"
"""

import json
import os
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime
import re
from collections import Counter

BASE_DIR = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
PROFILE_PATH = os.path.join(BASE_DIR, "data", "master_profile.json")
OUTPUT_DIR = os.path.join(BASE_DIR, "generated")

class CVGenerator:
    """Generate ATS-optimized CVs."""
    
    def __init__(self):
        with open(PROFILE_PATH, 'r', encoding='utf-8') as f:
            self.profile = json.load(f)
    
    def extract_keywords(self, job_description: str) -> list:
        """Extract relevant keywords from job description."""
        # Common project controls keywords
        keywords_db = [
            "primavera", "p6", "ms project", "scheduling", "planning",
            "epc", "oil gas", "offshore", "project controls", "claims",
            "eot", "extension of time", "fidic", "delay analysis",
            "earned value", "evm", "spi", "cpi", "cost control",
            "baseline", "critical path", "float", "recovery",
            "progress", "reporting", "dashboard", "power bi",
            "pmp", "pmc", "contractor", "risk", "governance"
        ]
        
        jd_lower = job_description.lower()
        found = []
        
        for kw in keywords_db:
            if kw in jd_lower:
                found.append(kw)
        
        # Also find profile keywords in JD
        for kw in self.profile.get("keywords", []):
            if kw.lower() in jd_lower:
                found.append(kw)
        
        return list(set(found))
    
    def calculate_ats_score(self, job_description: str) -> dict:
        """Calculate ATS match score."""
        jd_words = set(re.findall(r'\b\w+\b', job_description.lower()))
        cv_keywords = set(kw.lower() for kw in self.profile.get("keywords", []))
        
        matched = jd_words & cv_keywords
        
        # Calculate competency matches
        competencies = self.profile.get("core_competencies", [])
        comp_matches = sum(1 for c in competencies if any(w in job_description.lower() for w in c.lower().split()))
        
        return {
            "keyword_matches": len(matched),
            "total_keywords": len(cv_keywords),
            "competency_matches": comp_matches,
            "total_competencies": len(competencies),
            "score_percent": round((len(matched) / max(len(cv_keywords), 1)) * 100, 1),
            "matched_keywords": list(matched)
        }
    
    def generate_cv(self, job_title: str = None, company: str = None, 
                    job_description: str = None, output_name: str = None) -> str:
        """Generate tailored CV document."""
        
        doc = Document()
        
        # Set margins
        for section in doc.sections:
            section.top_margin = Inches(0.5)
            section.bottom_margin = Inches(0.5)
            section.left_margin = Inches(0.6)
            section.right_margin = Inches(0.6)
        
        p = self.profile["personal"]
        
        # === HEADER ===
        name = doc.add_paragraph()
        name_run = name.add_run(p["name"] + ", " + p["credentials"])
        name_run.bold = True
        name_run.font.size = Pt(16)
        name.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Title
        title = doc.add_paragraph()
        title_run = title.add_run(p["title"])
        title_run.bold = True
        title_run.font.size = Pt(11)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Industries
        industries = doc.add_paragraph()
        ind_run = industries.add_run("EPC • Oil & Gas • Offshore • Infrastructure")
        ind_run.font.size = Pt(10)
        industries.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Contact
        contact = doc.add_paragraph()
        contact_text = f"{p['location']} | {p['phone']} | {p['email_primary']}"
        contact.add_run(contact_text).font.size = Pt(10)
        contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # LinkedIn
        linkedin = doc.add_paragraph()
        linkedin.add_run(f"LinkedIn: {p['linkedin']}").font.size = Pt(10)
        linkedin.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # === PROFESSIONAL SUMMARY ===
        self._add_section_header(doc, "PROFESSIONAL SUMMARY")
        summary = doc.add_paragraph()
        summary.add_run(self.profile["professional_summary"]).font.size = Pt(10)
        
        # === CORE COMPETENCIES ===
        self._add_section_header(doc, "CORE COMPETENCIES")
        
        # Two column competencies
        comps = self.profile["core_competencies"]
        mid = len(comps) // 2
        
        for i in range(mid):
            line = doc.add_paragraph()
            left = comps[i] if i < len(comps) else ""
            right = comps[mid + i] if (mid + i) < len(comps) else ""
            line.add_run(f"• {left:<40} • {right}").font.size = Pt(10)
        
        # === LEADERSHIP & GOVERNANCE ===
        self._add_section_header(doc, "LEADERSHIP & GOVERNANCE")
        for item in self.profile["leadership_governance"]:
            p = doc.add_paragraph(style='List Bullet')
            p.add_run(item).font.size = Pt(10)
        
        # === PROFESSIONAL EXPERIENCE ===
        self._add_section_header(doc, "PROFESSIONAL EXPERIENCE")
        
        for exp in self.profile["experience"]:
            # Job title and company
            job = doc.add_paragraph()
            job.add_run(exp["title"]).bold = True
            job.add_run(f"\n{exp['company']} – {exp['industry']} | {exp['period']}")
            
            if "project_value" in exp:
                job.add_run(f" | {exp['project_value']}")
            
            # Highlights
            for h in exp["highlights"]:
                bullet = doc.add_paragraph(style='List Bullet')
                bullet.add_run(h).font.size = Pt(10)
        
        # === EDUCATION & CERTIFICATIONS ===
        self._add_section_header(doc, "EDUCATION & CERTIFICATIONS")
        
        for edu in self.profile["education"]:
            p = doc.add_paragraph()
            p.add_run(f"{edu['degree']} – {edu['institution']}").font.size = Pt(10)
        
        for cert in self.profile["certifications"]:
            p = doc.add_paragraph()
            p.add_run(cert).font.size = Pt(10)
        
        # === TECHNICAL SKILLS ===
        self._add_section_header(doc, "TECHNICAL SKILLS")
        skills = doc.add_paragraph()
        skills.add_run(" | ".join(self.profile["technical_skills"])).font.size = Pt(10)
        
        # === AVAILABILITY ===
        avail = doc.add_paragraph()
        personal = self.profile["personal"]
        avail.add_run(f"\n{personal['relocation']} | {personal['availability']} Availability").font.size = Pt(10)
        avail.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Save
        if not output_name:
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            if company:
                output_name = f"CV_TahirYamin_{company.replace(' ', '_')}_{timestamp}.docx"
            else:
                output_name = f"CV_TahirYamin_{timestamp}.docx"
        
        output_path = os.path.join(OUTPUT_DIR, output_name)
        os.makedirs(OUTPUT_DIR, exist_ok=True)
        doc.save(output_path)
        
        print(f"✓ Generated: {output_path}")
        
        if job_description:
            score = self.calculate_ats_score(job_description)
            print(f"✓ ATS Score: {score['score_percent']}%")
            print(f"  Keywords matched: {score['matched_keywords'][:10]}")
        
        return output_path
    
    def _add_section_header(self, doc, title: str):
        """Add section header with line."""
        p = doc.add_paragraph()
        p.add_run(title).bold = True
        p.add_run().font.size = Pt(11)


def main():
    import argparse
    
    parser = argparse.ArgumentParser(description="Generate ATS-Optimized CV")
    parser.add_argument("--job-file", type=str, help="Path to job description file")
    parser.add_argument("--company", type=str, help="Company name")
    parser.add_argument("--title", type=str, help="Job title")
    parser.add_argument("--score-only", action="store_true", help="Only calculate ATS score")
    
    args = parser.parse_args()
    
    gen = CVGenerator()
    
    job_desc = None
    if args.job_file and os.path.exists(args.job_file):
        with open(args.job_file, 'r', encoding='utf-8') as f:
            job_desc = f.read()
    
    if args.score_only and job_desc:
        score = gen.calculate_ats_score(job_desc)
        print(f"ATS Score: {score['score_percent']}%")
        print(f"Keyword matches: {score['keyword_matches']}/{score['total_keywords']}")
        print(f"Competency matches: {score['competency_matches']}/{score['total_competencies']}")
        print(f"Top keywords: {score['matched_keywords'][:15]}")
    else:
        gen.generate_cv(
            job_title=args.title,
            company=args.company,
            job_description=job_desc
        )


if __name__ == "__main__":
    main()
