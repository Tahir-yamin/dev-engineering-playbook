import json
import os
from datetime import datetime
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

BASE_DIR = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
PROFILE_PATH = os.path.join(BASE_DIR, "data", "master_profile.json")
OUTPUT_DIR = os.path.join(BASE_DIR, "generated")

class CoverLetterGenerator:
    """Generate modern, personalized cover letters in DOCX format."""
    
    def __init__(self):
        with open(PROFILE_PATH, 'r', encoding='utf-8') as f:
            self.profile = json.load(f)
    
    def generate(self, company: str, role: str, hiring_manager: str = None,
                 job_highlights: list = None, why_company: str = None) -> str:
        """Generate cover letter as DOCX."""
        
        p = self.profile["personal"]
        name = p["name"]
        
        # Default why company if not provided
        if not why_company:
            why_company = f"{company}'s reputation for excellence in EPC and project delivery"
        
        # Match top 3 qualifications
        top_quals = [
            f"15+ years leading project controls and planning for EPC, Oil & Gas, and offshore projects valued up to USD 750M",
            f"Expert in Primavera P6 (L1-L4), schedule risk analysis, delay analysis, and EOT/claims support under FIDIC contracts",
            f"Proven leadership managing planning teams, establishing governance frameworks, and delivering commercially focused reporting"
        ]
        
        # Create DOCX
        doc = Document()
        
        # Set margins
        for section in doc.sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)
        
        # Header - Contact Info
        header = doc.add_paragraph()
        header.add_run(name).bold = True
        header.add_run(f"\n{p['location']} | {p['phone']} | {p['email_primary']}")
        header.add_run(f"\nLinkedIn: {p['linkedin']}")
        
        # Date
        date_p = doc.add_paragraph()
        date_p.add_run(f"\n{datetime.now().strftime('%B %d, %Y')}")
        
        # Recipient
        recipient = doc.add_paragraph()
        recipient.add_run(f"\n{'Hiring Manager' if not hiring_manager else hiring_manager}")
        recipient.add_run(f"\n{company}")
        
        # Subject
        subject = doc.add_paragraph()
        subject.add_run(f"\nRE: Application for {role}").bold = True
        
        # Salutation
        doc.add_paragraph(f"\nDear {'Hiring Manager' if not hiring_manager else hiring_manager},")
        
        # Opening paragraph
        opening = doc.add_paragraph()
        opening.add_run(f"I am writing to express my strong interest in the {role} position at {company}. With 15+ years of experience in project controls, planning, and schedule risk management across EPC, Oil & Gas, offshore, and infrastructure projects, I am confident I can contribute immediately to your team's success.")
        
        # Qualifications
        doc.add_paragraph("\nWhat makes me an excellent fit for this role:")
        
        for qual in top_quals:
            bullet = doc.add_paragraph(style='List Bullet')
            bullet.add_run(qual)
        
        # Company interest
        interest = doc.add_paragraph()
        interest.add_run(f"\nI am particularly drawn to {why_company}. My experience working with multinational EPCs and PMCs, combined with my PMP certification and hands-on expertise in claims support and commercial planning, positions me to add value from day one.")
        
        # Experience highlight
        exp = doc.add_paragraph()
        exp.add_run("\nIn my current role as Senior Planning Engineer at Pakistan State Oil, I lead project planning and controls for nationwide infrastructure projects, developing Level-3/L4 baselines, performing EVM analysis, and supporting EOT evaluations under FIDIC principles. Previously at Karachi Shipyard, I managed a 15-engineer team controlling USD 750M+ in EPC projects.")
        
        # Call to action
        cta = doc.add_paragraph()
        cta.add_run(f"\nI would welcome the opportunity to discuss how my skills in project controls, delay analysis, and programme governance can support {company}'s project delivery objectives.")
        
        # Closing
        doc.add_paragraph("\nThank you for considering my application. I look forward to speaking with you.")
        
        # Signature
        sig = doc.add_paragraph()
        sig.add_run("\nBest regards,")
        sig.add_run(f"\n\n{name}, PMP")
        sig.add_run(f"\n{p['phone']}")
        sig.add_run(f"\n{p['email_primary']}")
        
        # Save as DOCX
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        filename = f"CoverLetter_{company.replace(' ', '_')}_{role.replace(' ', '_')}_{timestamp}.docx"
        output_path = os.path.join(OUTPUT_DIR, filename)
        os.makedirs(OUTPUT_DIR, exist_ok=True)
        
        doc.save(output_path)
        
        print(f"✓ Generated: {output_path}")
        return output_path


def main():
    import argparse
    
    parser = argparse.ArgumentParser(description="Generate Cover Letter")
    parser.add_argument("--company", type=str, required=True, help="Company name")
    parser.add_argument("--role", type=str, required=True, help="Job title")
    parser.add_argument("--manager", type=str, help="Hiring manager name")
    parser.add_argument("--why", type=str, help="Why interested in company")
    
    args = parser.parse_args()
    
    gen = CoverLetterGenerator()
    gen.generate(
        company=args.company,
        role=args.role,
        hiring_manager=args.manager,
        why_company=args.why
    )


if __name__ == "__main__":
    main()
